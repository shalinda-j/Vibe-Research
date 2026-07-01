"""Export a Markdown research report to PDF.

Optional feature (the ``[pdf]`` extra → ``fpdf2``). The report Markdown is
rendered to HTML with markdown-it-py — already present as a Textual dependency,
so the HTML step needs nothing new — and laid out by fpdf2 using a *discovered
system Unicode font*. That means accents, arrows, bullets, em-dashes and the
like render correctly without shipping a font file or pulling in any native
dependency (no GTK/Cairo/Pango), which keeps it painless on Windows.

Everything heavy is imported lazily inside the functions, so ``import
vibe_research.export`` stays cheap and ``pdf_available()`` can be probed without
fpdf2 installed.
"""

from __future__ import annotations

import importlib.util
import re
import sys
from datetime import datetime
from pathlib import Path

# (regular, bold, italic, bold-italic) TrueType families to try, in order.
_FONT_FAMILIES = (
    ("arial.ttf", "arialbd.ttf", "ariali.ttf", "arialbi.ttf"),
    ("segoeui.ttf", "segoeuib.ttf", "segoeuii.ttf", "segoeuiz.ttf"),
    ("DejaVuSans.ttf", "DejaVuSans-Bold.ttf", "DejaVuSans-Oblique.ttf", "DejaVuSans-BoldOblique.ttf"),
    ("Helvetica.ttc", "Helvetica.ttc", "Helvetica.ttc", "Helvetica.ttc"),
)


def _font_dirs() -> list[Path]:
    if sys.platform == "win32":
        win = Path(__import__("os").environ.get("WINDIR", r"C:\Windows"))
        return [win / "Fonts", Path.home() / "AppData/Local/Microsoft/Windows/Fonts"]
    if sys.platform == "darwin":
        return [Path("/System/Library/Fonts"), Path("/Library/Fonts"), Path.home() / "Library/Fonts"]
    return [
        Path("/usr/share/fonts"),
        Path("/usr/local/share/fonts"),
        Path.home() / ".fonts",
        Path.home() / ".local/share/fonts",
    ]


def _find_unicode_font() -> tuple[Path, Path, Path, Path] | None:
    """Locate a regular/bold/italic/bold-italic Unicode TTF set on this machine."""
    dirs = _font_dirs()
    # Fast path: exact filenames in a known font directory.
    for family in _FONT_FAMILIES:
        for directory in dirs:
            paths = [directory / name for name in family]
            if all(p.exists() for p in paths):
                return paths[0], paths[1], paths[2], paths[3]
    # Slow path: recursive search for DejaVuSans (bundled with many Linux distros).
    for directory in dirs:
        if not directory.exists():
            continue
        try:
            hits = {p.name: p for p in directory.rglob("DejaVuSans*.ttf")}
        except OSError:
            continue
        reg = hits.get("DejaVuSans.ttf")
        if reg:
            return (
                reg,
                hits.get("DejaVuSans-Bold.ttf", reg),
                hits.get("DejaVuSans-Oblique.ttf", reg),
                hits.get("DejaVuSans-BoldOblique.ttf", reg),
            )
    return None


_IMG_REF = re.compile(r"(!\[[^\]]*\]\()([^)]+)(\))")
_IMG_REF_ALT = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def _pdf_images(markdown: str, base_dir) -> str:
    """Prepare image refs for fpdf2, which loads every <img> eagerly and aborts
    the whole PDF on any failure. Remote/data images are dropped to their caption
    (fpdf2 would fetch them and a 404/timeout kills the render), and local images
    are embedded only if the file actually exists — else replaced with the caption.
    """
    base = Path(base_dir) if base_dir else None

    def _sub(match):
        alt, src = match.group(1), match.group(2).strip()
        if src.startswith(("http://", "https://", "data:")):
            return f"*{alt}*" if alt else ""
        path = Path(src)
        if not path.is_absolute() and base is not None:
            path = (base / src).resolve()
        if path.is_file():
            return f"![{alt}]({path.as_posix()})"
        return f"*{alt}*" if alt else ""

    return _IMG_REF_ALT.sub(_sub, markdown or "")


def _resolve_images(markdown: str, base_dir) -> str:
    """Rewrite relative image paths to absolute (against base_dir) so PDF/DOCX
    renderers — which don't know the report's directory — can find local charts."""
    if not base_dir:
        return markdown
    base = Path(base_dir)

    def _sub(match):
        src = match.group(2).strip()
        if src.startswith(("http://", "https://", "data:")) or Path(src).is_absolute():
            return match.group(0)
        # POSIX-style path: fpdf2's write_html percent-encodes backslashes.
        return match.group(1) + (base / src).resolve().as_posix() + match.group(3)

    return _IMG_REF.sub(_sub, markdown or "")


def pdf_available() -> bool:
    """True if the PDF engine (fpdf2) is importable — probe without importing it."""
    return importlib.util.find_spec("fpdf") is not None


def _latin1_safe(text: str) -> str:
    """Drop-in for the no-Unicode-font fallback: keep only Latin-1 renderable chars."""
    return text.encode("latin-1", "replace").decode("latin-1")


def markdown_to_pdf(markdown: str, out_path: Path | str, title: str = "", base_dir=None) -> Path:
    """Render a Markdown report to a PDF file and return its path.

    Raises ``RuntimeError`` with install instructions if fpdf2 is missing.
    """
    try:
        from fpdf import FPDF
    except ImportError as exc:  # pragma: no cover - exercised only without the extra
        raise RuntimeError(
            "PDF export needs the 'fpdf2' package.\n"
            "    pip install fpdf2\n"
            '    (or: pip install "vibe-research[pdf]")'
        ) from exc
    from markdown_it import MarkdownIt

    markdown = _pdf_images(markdown, base_dir)
    html = MarkdownIt("commonmark", {"html": False}).enable("table").render(markdown or "")
    fonts = _find_unicode_font()

    class _ReportPDF(FPDF):
        def footer(self) -> None:
            self.set_y(-12)
            self.set_font(self._body_family, "I", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 8, f"vibe-research  ·  page {self.page_no()}/{{nb}}", align="C")
            self.set_text_color(0, 0, 0)

    pdf = _ReportPDF(format="A4")
    pdf._body_family = "helvetica"
    if fonts:
        reg, bold, ital, bolditalic = fonts
        pdf.add_font("report", "", str(reg))
        pdf.add_font("report", "B", str(bold))
        pdf.add_font("report", "I", str(ital))
        pdf.add_font("report", "BI", str(bolditalic))
        pdf._body_family = "report"

    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(18, 16, 18)
    pdf.add_page()

    # --- title block ---------------------------------------------------------
    if title:
        pdf.set_font(pdf._body_family, "B", 19)
        pdf.multi_cell(0, 9, title if fonts else _latin1_safe(title))
        pdf.ln(1)
    pdf.set_font(pdf._body_family, "I", 9)
    pdf.set_text_color(120, 120, 120)
    pdf.multi_cell(0, 6, f"Generated {datetime.now():%Y-%m-%d %H:%M} by vibe-research")
    pdf.set_draw_color(210, 210, 210)
    pdf.line(pdf.l_margin, pdf.get_y() + 1, pdf.w - pdf.r_margin, pdf.get_y() + 1)
    pdf.ln(4)
    pdf.set_text_color(0, 0, 0)

    # --- body ----------------------------------------------------------------
    pdf.set_font(pdf._body_family, "", 11)
    pdf.write_html(html if fonts else _latin1_safe(html))

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(out_path))
    return out_path


def pdf_path_for(report_path: Path | str) -> Path:
    """The sibling ``.pdf`` path for a saved ``.md`` report."""
    return Path(report_path).with_suffix(".pdf")


def html_path_for(report_path: Path | str) -> Path:
    """The sibling ``.html`` path for a saved ``.md`` report."""
    return Path(report_path).with_suffix(".html")


_HTML_TEMPLATE = """\
<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
<style>
  :root {{ color-scheme: light dark; }}
  body {{
    max-width: 46rem; margin: 2.5rem auto; padding: 0 1.2rem;
    font: 16px/1.65 -apple-system, Segoe UI, Roboto, Helvetica, Arial, sans-serif;
    color: #1a1a1a; background: #fff;
  }}
  h1, h2, h3 {{ line-height: 1.25; margin-top: 1.8em; }}
  h1 {{ font-size: 1.9rem; border-bottom: 2px solid #eee; padding-bottom: .3em; }}
  h2 {{ font-size: 1.4rem; }}
  a {{ color: #2563eb; word-break: break-word; }}
  code {{ background: #f3f3f3; padding: .1em .35em; border-radius: 4px; font-size: .9em; }}
  pre {{ background: #f6f8fa; padding: 1em; border-radius: 8px; overflow-x: auto; }}
  pre code {{ background: none; padding: 0; }}
  blockquote {{ margin: 1em 0; padding: .2em 1em; border-left: 4px solid #d1d5db; color: #555; }}
  table {{ border-collapse: collapse; width: 100%; }}
  th, td {{ border: 1px solid #ddd; padding: .5em .7em; text-align: left; }}
  th {{ background: #f6f8fa; }}
  hr {{ border: none; border-top: 1px solid #eee; margin: 2em 0; }}
  img {{ max-width: 100%; height: auto; display: block; margin: 1.2em auto; border-radius: 6px; }}
  .mermaid {{ text-align: center; margin: 1.6em 0; }}
  footer {{ margin-top: 3rem; padding-top: 1rem; border-top: 1px solid #eee;
            color: #888; font-size: .85rem; }}
  @media (prefers-color-scheme: dark) {{
    body {{ color: #e6e6e6; background: #16181d; }}
    h1 {{ border-color: #2a2d34; }} code {{ background: #23262e; }}
    pre {{ background: #1c1f26; }} th {{ background: #23262e; }}
    td, th {{ border-color: #2a2d34; }} a {{ color: #6ea8fe; }}
  }}
</style>
</head>
<body>
{body}
<footer>Generated by vibe-research.</footer>
</body>
</html>
"""


_MERMAID_CODE = re.compile(
    r'<pre><code class="language-mermaid">(.*?)</code></pre>', re.DOTALL
)
_MERMAID_CDN = (
    '<script type="module">'
    'import mermaid from "https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.esm.min.mjs";'
    'mermaid.initialize({ startOnLoad: true, theme: "neutral" });'
    "</script>"
)


def _mermaidify(html: str) -> tuple[str, bool]:
    """Turn ```mermaid code blocks into mermaid.js <div>s. Returns (html, had_any).

    The block content is left HTML-*escaped*: the browser decodes it into the
    element's textContent (which is what mermaid.js parses, so arrows like ``-->``
    still work), while any escaped ``<script>`` stays inert text rather than a
    live node — no XSS from untrusted diagram content.
    """
    found = {"any": False}

    def _sub(match):
        found["any"] = True
        return f'<div class="mermaid">{match.group(1)}</div>'

    return _MERMAID_CODE.sub(_sub, html), found["any"]


def markdown_to_html(markdown: str, title: str = "") -> str:
    """Render a Markdown report into a standalone, styled HTML document string.

    ```mermaid` blocks become live diagrams (rendered by mermaid.js), and image
    references (rendered charts, figures) become responsive ``<img>`` tags.
    """
    from markdown_it import MarkdownIt

    # html=False so raw HTML in LLM/web-derived report text is escaped, not
    # injected live into the exported page.
    body = MarkdownIt("commonmark", {"html": False}).enable("table").render(markdown or "")
    body, has_mermaid = _mermaidify(body)
    if has_mermaid:
        body += "\n" + _MERMAID_CDN
    safe_title = (title or "vibe-research report").replace("<", "&lt;").replace(">", "&gt;")
    return _HTML_TEMPLATE.format(title=safe_title, body=body)


def markdown_to_html_file(markdown: str, out_path: Path | str, title: str = "") -> Path:
    """Render a Markdown report to an HTML file and return its path."""
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(markdown_to_html(markdown, title), encoding="utf-8")
    return out_path


# --------------------------------------------------------------------- DOCX


def docx_available() -> bool:
    """True if the DOCX engine (python-docx) is importable."""
    return importlib.util.find_spec("docx") is not None


def docx_path_for(report_path: Path | str) -> Path:
    """The sibling ``.docx`` path for a saved ``.md`` report."""
    return Path(report_path).with_suffix(".docx")


def _emit_docx_inline(paragraph, inline) -> None:
    """Render one markdown-it inline token's children into a docx paragraph."""
    bold = italic = False
    href = None
    for child in (inline.children or []):
        kind = child.type
        if kind == "strong_open":
            bold = True
        elif kind == "strong_close":
            bold = False
        elif kind == "em_open":
            italic = True
        elif kind == "em_close":
            italic = False
        elif kind == "link_open":
            href = child.attrGet("href")
        elif kind == "link_close":
            if href:
                run = paragraph.add_run(f" ({href})")
                run.italic = True
            href = None
        elif kind in ("softbreak", "hardbreak"):
            paragraph.add_run("\n")
        elif kind == "code_inline":
            run = paragraph.add_run(child.content)
            run.font.name = "Consolas"
        elif kind == "image":
            src = child.attrGet("src") or ""
            try:
                if src and not src.startswith(("http://", "https://", "data:")) and Path(src).is_file():
                    from docx.shared import Inches

                    paragraph.add_run().add_picture(str(src), width=Inches(6.0))
                else:  # remote/missing image -> show its caption as a placeholder
                    run = paragraph.add_run(f"[{child.content or 'image'}]")
                    run.italic = True
            except Exception:
                pass
        elif kind == "text":
            run = paragraph.add_run(child.content)
            run.bold = bold
            run.italic = italic


def _consume_docx_table(document, tokens, start: int) -> int:
    """Render a markdown-it table (from ``table_open`` at ``start``) as a Word
    table. Returns the index of the matching ``table_close``."""
    rows: list[list] = []
    current: list | None = None
    i = start + 1
    while i < len(tokens) and tokens[i].type != "table_close":
        kind = tokens[i].type
        if kind == "tr_open":
            current = []
        elif kind == "tr_close":
            if current is not None:
                rows.append(current)
                current = None
        elif kind in ("th_open", "td_open"):
            if current is not None:
                current.append(tokens[i + 1] if i + 1 < len(tokens) else None)
            i += 1  # skip the cell's inline token
        i += 1
    if rows:
        ncols = max(len(r) for r in rows)
        table = document.add_table(rows=len(rows), cols=ncols)
        try:
            table.style = "Light Grid Accent 1"
        except Exception:
            pass
        for r, cells in enumerate(rows):
            for c, inline in enumerate(cells):
                if inline is not None:
                    _emit_docx_inline(table.cell(r, c).paragraphs[0], inline)
    return i


def markdown_to_docx(markdown: str, out_path: Path | str, title: str = "", base_dir=None) -> Path:
    """Render a Markdown report to a Word ``.docx`` file and return its path.

    Handles headings (with inline formatting), paragraphs, bold/italic, links
    (text + URL), ordered and bullet lists, code blocks, tables (rendered as Word
    tables), and embedded local images (e.g. rendered charts). Raises
    ``RuntimeError`` with install instructions if python-docx is missing.
    """
    try:
        import docx
        from docx.shared import Pt
    except ImportError as exc:  # pragma: no cover - exercised only without the extra
        raise RuntimeError(
            "DOCX export needs the 'python-docx' package.\n"
            "    pip install python-docx\n"
            '    (or: pip install "vibe-research[docx]")'
        ) from exc
    from markdown_it import MarkdownIt

    markdown = _resolve_images(markdown, base_dir)
    document = docx.Document()
    if title:
        document.add_heading(title, level=0)

    tokens = MarkdownIt("commonmark").enable("table").parse(markdown or "")
    list_stack: list[str] = []
    idx, n = 0, len(tokens)
    while idx < n:
        tok = tokens[idx]
        kind = tok.type
        if kind == "heading_open":
            level = int(tok.tag[1]) if tok.tag[1:].isdigit() else 2
            heading = document.add_heading("", level=min(level, 4))
            _emit_docx_inline(heading, tokens[idx + 1])  # render inline markup, not raw source
            idx += 3
            continue
        if kind == "table_open":
            idx = _consume_docx_table(document, tokens, idx) + 1
            continue
        if kind == "paragraph_open":
            if list_stack:
                style = "List Bullet" if list_stack[-1] == "ul" else "List Number"
                paragraph = document.add_paragraph(style=style)
            else:
                paragraph = document.add_paragraph()
            _emit_docx_inline(paragraph, tokens[idx + 1])
            idx += 3
            continue
        if kind == "bullet_list_open":
            list_stack.append("ul")
        elif kind == "ordered_list_open":
            list_stack.append("ol")
        elif kind in ("bullet_list_close", "ordered_list_close"):
            if list_stack:
                list_stack.pop()
        elif kind in ("fence", "code_block"):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(tok.content.rstrip("\n"))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        idx += 1

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    return out_path
