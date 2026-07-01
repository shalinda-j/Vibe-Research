"""Offline tests for vibe-research.

Runnable with no network and no heavy deps:
    PYTHONPATH=src python -m unittest discover -s tests -v

The orchestrator tests drive the whole multi-agent loop
(plan -> research -> verify/debate -> critique -> [re-plan] -> write) against a
FakeBackend, validating orchestration, data-validation, the self-refining
iteration, and long-term memory — all without any API calls.
"""

import asyncio
import importlib.util
import os
import pathlib
import sys
import tempfile
import unittest

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parents[1] / "src"))

from vibe_research.backends import Backend, choose_mode, extract_urls  # noqa: E402
from vibe_research import config as cfgmod  # noqa: E402
from vibe_research import enrich  # noqa: E402
from vibe_research import export as exportmod  # noqa: E402
from vibe_research import visuals  # noqa: E402
from vibe_research import reports as reportsmod  # noqa: E402
from vibe_research import schemas  # noqa: E402
from vibe_research.agents import (  # noqa: E402
    PlannerAgent,
    ResearcherAgent,
    VerifierAgent,
    _extract_confidence,
)
from vibe_research.memory import Memory  # noqa: E402
from vibe_research.pipeline import run_pipeline  # noqa: E402
from vibe_research.schemas import (  # noqa: E402
    Critique,
    Finding,
    ResearchPlan,
    ResearchRecord,
    ValidationError,
    VerificationReport,
    aggregate_verifications,
)


# --------------------------------------------------------------------- fakes


class FakeBackend(Backend):
    """Canned responses keyed off distinctive prompt substrings, one per role."""

    name = "fake"

    def __init__(self):
        self.calls = []

    async def _complete_once(self, prompt, model, use_search=False):  # satisfies the ABC
        raise NotImplementedError

    async def complete(self, prompt, model, use_search=False):
        self.calls.append((model, use_search))
        if "research lead" in prompt and "Return ONLY a JSON object" in prompt:
            return (
                '{"subquestions": ['
                '{"text": "Question one about X?", "rationale": "a"},'
                '{"text": "Question two about Y?", "rationale": "b"},'
                '{"text": "Question three about Z?", "rationale": "c"}]}'
            ), []
        if "refining coverage of" in prompt:
            return '{"subquestions": [{"text": "A deeper follow-up about W?"}]}', []
        if "skeptical fact-checker" in prompt:
            return (
                '{"verdicts": [{"claim": "X is Y", "status": "supported", '
                '"confidence": 0.9, "note": "cited"}], "gaps": [], '
                '"contradictions": [], "overall_confidence": 0.9}'
            ), []
        if "managing editor" in prompt:
            return '{"approved": true, "quality": 0.9, "missing": [], "issues": [], "note": "solid"}', []
        if "knowledgeable human writer" in prompt:
            return (
                "# Findings\n\nA cleaner, human-sounding version grounded in "
                "https://example.com/a.\n\n## Confidence & Gaps\nSolid overall; pricing uncertain."
            ), []
        if "research report in Markdown" in prompt:
            return (
                "# Findings\n\nBody grounded in https://example.com/a\n\n"
                "## Confidence & Gaps\nSolid overall; pricing uncertain."
            ), []
        # researcher (web-search) prompt
        return (
            "Answer citing https://example.com/x and https://example.com/y.\nCONFIDENCE: 0.9",
            ["https://example.com/x", "https://example.com/y"],
        )


class IteratingBackend(FakeBackend):
    """Like FakeBackend but the editor rejects once, forcing one gap round."""

    def __init__(self):
        super().__init__()
        self.editor_calls = 0

    async def complete(self, prompt, model, use_search=False):
        if "managing editor" in prompt:
            self.editor_calls += 1
            self.calls.append((model, use_search))
            if self.editor_calls == 1:
                return (
                    '{"approved": false, "quality": 0.4, '
                    '"missing": ["A deeper follow-up about W?"], "note": "need more"}'
                ), []
            return '{"approved": true, "quality": 0.9, "missing": [], "note": "ok"}', []
        return await super().complete(prompt, model, use_search)


class FailResearchBackend(FakeBackend):
    """Every web-search (researcher) call raises — simulates a rate limit."""

    async def complete(self, prompt, model, use_search=False):
        if use_search:
            raise RuntimeError("simulated rate limit")
        return await super().complete(prompt, model, use_search)


class FailVerifyBackend(FakeBackend):
    """Every fact-check call raises — simulates a transient backend timeout."""

    async def complete(self, prompt, model, use_search=False):
        if "skeptical fact-checker" in prompt:
            raise RuntimeError("simulated timeout")
        return await super().complete(prompt, model, use_search)


class StripUrlBackend(FakeBackend):
    """The humanizer returns prose with no URLs — should trigger the fallback."""

    async def complete(self, prompt, model, use_search=False):
        if "knowledgeable human writer" in prompt:
            return "Rewritten with every link removed.\n\n## Confidence & Gaps\nok", []
        return await super().complete(prompt, model, use_search)


class DropSectionBackend(FakeBackend):
    """Humanizer keeps the URL but drops 'Confidence & Gaps' — must fall back."""

    async def complete(self, prompt, model, use_search=False):
        if "knowledgeable human writer" in prompt:
            return "Rewritten prose still citing https://example.com/a, no honesty section.", []
        return await super().complete(prompt, model, use_search)


# --- concrete backends that exercise the base reliability wrapper -----------


class FlakyBackend(Backend):
    """Fails a set number of times, then succeeds — to test retry/backoff."""

    name = "flaky"

    def __init__(self, fail_times=2):
        super().__init__()
        self.base_delay = 0.0     # no real sleep in tests
        self.attempts = 0
        self.fail_times = fail_times

    async def _complete_once(self, prompt, model, use_search=False):
        self.attempts += 1
        if self.attempts <= self.fail_times:
            raise ConnectionError("transient")
        return "ok", []


class AuthErrorBackend(Backend):
    """Raises a non-retryable 401 — must fail fast without retrying."""

    name = "auth"

    def __init__(self):
        super().__init__()
        self.base_delay = 0.0
        self.attempts = 0

    async def _complete_once(self, prompt, model, use_search=False):
        self.attempts += 1
        exc = RuntimeError("bad key")
        exc.status_code = 401
        raise exc


class HangBackend(Backend):
    """Sleeps forever — to test the per-call timeout."""

    name = "hang"

    def __init__(self):
        super().__init__()
        self.max_retries = 0
        self.call_timeout = 0.05

    async def _complete_once(self, prompt, model, use_search=False):
        await asyncio.sleep(5)
        return "never", []


class ContradictionBackend(FakeBackend):
    """Fact-checker reports a contradiction — to test the Disagreements section."""

    async def complete(self, prompt, model, use_search=False):
        if "skeptical fact-checker" in prompt:
            return (
                '{"verdicts":[{"claim":"X","status":"contradicted","confidence":0.4}],'
                '"gaps":[],"contradictions":["Source A says X; Source B says not-X"],'
                '"overall_confidence":0.5}'
            ), []
        return await super().complete(prompt, model, use_search)


class BlockedSourceBackend(FakeBackend):
    """Researcher cites a blocked + a kept domain — to test source filtering."""

    async def complete(self, prompt, model, use_search=False):
        if use_search:
            return (
                "Answer citing https://reddit.com/r/x and https://cdc.gov/y.\nCONFIDENCE: 0.8",
                ["https://reddit.com/r/x", "https://cdc.gov/y"],
            )
        return await super().complete(prompt, model, use_search)


class ThrottleBackend(Backend):
    """Records peak concurrency — to test the max_concurrency cap."""

    name = "throttle"

    def __init__(self):
        super().__init__()
        self.configure(max_concurrency=2)
        self.active = 0
        self.max_active = 0

    async def _complete_once(self, prompt, model, use_search=False):
        self.active += 1
        self.max_active = max(self.max_active, self.active)
        await asyncio.sleep(0.02)
        self.active -= 1
        return "ok", []


# --------------------------------------------------------------------- units


class TestUtils(unittest.TestCase):
    def test_slug(self):
        self.assertEqual(reportsmod.slug("Hello, World!"), "hello-world")
        self.assertEqual(reportsmod.slug("   "), "research")
        self.assertLessEqual(len(reportsmod.slug("x" * 200)), 60)

    def test_extract_urls_dedup_and_trim(self):
        text = "see https://a.com/1. and https://a.com/1) plus https://b.com/2,"
        self.assertEqual(extract_urls(text), ["https://a.com/1", "https://b.com/2"])

    def test_extract_urls_empty(self):
        self.assertEqual(extract_urls(""), [])
        self.assertEqual(extract_urls(None), [])


class TestSchemas(unittest.TestCase):
    def test_parse_json_plain_and_fenced(self):
        self.assertEqual(schemas.parse_json('{"a": 1}'), {"a": 1})
        self.assertEqual(schemas.parse_json('```json\n{"a": 1}\n```'), {"a": 1})

    def test_parse_json_salvages_chatty_output(self):
        self.assertEqual(
            schemas.parse_json('Sure! Here you go: {"a": [1, 2]} — hope that helps'),
            {"a": [1, 2]},
        )

    def test_parse_json_raises_on_garbage(self):
        with self.assertRaises(ValidationError):
            schemas.parse_json("no json at all")
        with self.assertRaises(ValidationError):
            schemas.parse_json("")

    def test_parse_json_salvages_first_of_multiple_blocks(self):
        # Regression: outermost-braces salvage used to fail on multiple blocks.
        self.assertEqual(
            schemas.parse_json('Here is one {"a": 1} and another {"b": 2}'),
            {"a": 1},
        )
        self.assertEqual(schemas.parse_json('Use {curly} braces: {"ok": true}'), {"ok": True})

    def test_parse_of_nonlist_fields_does_not_crash(self):
        # Regression: a scalar where a list is expected used to raise TypeError
        # that escaped the ValidationError safety net.
        report = VerificationReport.parse('{"verdicts": 5, "gaps": 42, "overall_confidence": 0.3}')
        self.assertEqual(report.verdicts, [])
        self.assertEqual(report.gaps, [])
        self.assertEqual(report.overall_confidence, 0.3)
        crit = Critique.parse('{"missing": 7, "issues": "oops", "approved": false}')
        self.assertEqual(crit.missing, [])
        self.assertEqual(crit.issues, [])  # a bare string is NOT iterated char-by-char

    def test_clamp01(self):
        self.assertEqual(schemas.clamp01(1.5), 1.0)
        self.assertEqual(schemas.clamp01(-2), 0.0)
        self.assertEqual(schemas.clamp01("0.4"), 0.4)
        self.assertEqual(schemas.clamp01("nonsense", default=0.3), 0.3)

    def test_clean_url(self):
        self.assertEqual(schemas.clean_url("https://x.com/a."), "https://x.com/a")
        self.assertEqual(schemas.clean_url("http://y.org/p?q=1"), "http://y.org/p?q=1")
        self.assertIsNone(schemas.clean_url("ftp://z.com"))
        self.assertIsNone(schemas.clean_url("not a url"))
        self.assertIsNone(schemas.clean_url(None))
        # dotted host now required (matches the documented contract)
        self.assertIsNone(schemas.clean_url("http://localhost"))
        self.assertEqual(schemas.clean_url("https://sub.example.co.uk/p"), "https://sub.example.co.uk/p")

    def test_clean_url_list_dedup_and_dicts(self):
        got = schemas.clean_url_list(
            ["https://a.com", {"url": "https://b.com"}, "https://a.com", "bad"]
        )
        self.assertEqual(got, ["https://a.com", "https://b.com"])

    def test_research_plan_parse_dedup_and_cap(self):
        raw = (
            '{"subquestions": ["What is X?", "What is X?", "How does Y work?", '
            '"too", "Another good question here?"]}'
        )
        plan = ResearchPlan.parse("topic", raw, hard_cap=2)
        # "too" is dropped (too short); dupes collapsed; capped to 2.
        self.assertEqual(len(plan.subquestions), 2)
        self.assertEqual(plan.questions[0], "What is X?")

    def test_research_plan_parse_raises_when_empty(self):
        with self.assertRaises(ValidationError):
            ResearchPlan.parse("topic", '{"subquestions": ["no", "x"]}')

    def test_verification_report_parse_and_weak_claims(self):
        raw = (
            '{"verdicts": [{"claim": "a", "status": "supported", "confidence": 0.9},'
            '{"claim": "b", "status": "unsupported", "confidence": 0.2}],'
            '"gaps": ["pricing"], "overall_confidence": 0.55}'
        )
        report = VerificationReport.parse(raw)
        self.assertEqual(len(report.verdicts), 2)
        self.assertEqual(report.overall_confidence, 0.55)
        self.assertEqual([v.claim for v in report.weak_claims], ["b"])

    def test_verification_status_normalisation(self):
        report = VerificationReport.parse(
            '{"verdicts": [{"claim": "a", "status": "TRUE"}]}'
        )
        self.assertEqual(report.verdicts[0].status, "supported")

    def test_aggregate_verifications_votes(self):
        r1 = VerificationReport.parse('{"gaps": ["g1"], "overall_confidence": 0.8}')
        r2 = VerificationReport.parse('{"gaps": ["g2", "g1"], "overall_confidence": 0.6}')
        agg = aggregate_verifications([r1, r2])
        self.assertEqual(agg.overall_confidence, 0.7)          # averaged vote
        self.assertEqual(sorted(agg.gaps), ["g1", "g2"])        # unioned, deduped

    def test_aggregate_verifications_empty_is_neutral(self):
        self.assertEqual(aggregate_verifications([]).overall_confidence, 0.5)

    def test_critique_parse(self):
        crit = Critique.parse('{"approved": false, "quality": 0.4, "missing": ["z"]}')
        self.assertFalse(crit.approved)
        self.assertEqual(crit.missing, ["z"])
        self.assertEqual(crit.quality, 0.4)


class TestConfig(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self._old = os.environ.get("XDG_CONFIG_HOME")
        os.environ["XDG_CONFIG_HOME"] = self._tmp.name

    def tearDown(self):
        if self._old is None:
            os.environ.pop("XDG_CONFIG_HOME", None)
        else:
            os.environ["XDG_CONFIG_HOME"] = self._old
        self._tmp.cleanup()

    def test_defaults_when_missing(self):
        cfg = cfgmod.load_config()
        self.assertEqual(cfg.mode, "auto")
        self.assertEqual(cfg.max_parallel, 2)
        self.assertEqual(cfg.max_iterations, 2)
        self.assertTrue(cfg.enable_memory)

    def test_roundtrip(self):
        cfg = cfgmod.default_config()
        cfg.mode = "subscription"
        cfg.max_parallel = 4
        cfg.enable_debate = False
        cfgmod.save_config(cfg)
        loaded = cfgmod.load_config()
        self.assertEqual(loaded.mode, "subscription")
        self.assertEqual(loaded.max_parallel, 4)
        self.assertFalse(loaded.enable_debate)

    def test_apply_setting_types_and_validation(self):
        cfg = cfgmod.default_config()
        cfgmod.apply_setting(cfg, "max_parallel", "5")
        self.assertEqual(cfg.max_parallel, 5)
        cfgmod.apply_setting(cfg, "verifier_votes", "3")
        self.assertEqual(cfg.verifier_votes, 3)
        cfgmod.apply_setting(cfg, "quality_threshold", "0.6")
        self.assertEqual(cfg.quality_threshold, 0.6)
        cfgmod.apply_setting(cfg, "enable_memory", "false")
        self.assertFalse(cfg.enable_memory)
        cfgmod.apply_setting(cfg, "enable_debate", "yes")
        self.assertTrue(cfg.enable_debate)
        cfgmod.apply_setting(cfg, "export_pdf", "true")
        self.assertTrue(cfg.export_pdf)

        # reliability knobs: 0 is allowed for retries/timeout (means "off")
        cfgmod.apply_setting(cfg, "max_retries", "0")
        self.assertEqual(cfg.max_retries, 0)
        cfgmod.apply_setting(cfg, "call_timeout", "0")
        self.assertEqual(cfg.call_timeout, 0)
        cfgmod.apply_setting(cfg, "max_concurrency", "3")
        self.assertEqual(cfg.max_concurrency, 3)

        # v0.4 sourcing / output knobs
        cfgmod.apply_setting(cfg, "citations", "plain")
        self.assertEqual(cfg.citations, "plain")
        cfgmod.apply_setting(cfg, "since_year", "0")            # 0 allowed (off)
        cfgmod.apply_setting(cfg, "since_year", "2020")
        self.assertEqual(cfg.since_year, 2020)
        cfgmod.apply_setting(cfg, "export_docx", "yes")
        self.assertTrue(cfg.export_docx)
        cfgmod.apply_setting(cfg, "verifier_model", "some-model")
        self.assertEqual(cfg.verifier_model, "some-model")
        cfgmod.apply_setting(cfg, "mode", "openai")
        self.assertEqual(cfg.mode, "openai")

        # v0.5 writing/visuals knobs
        cfgmod.apply_setting(cfg, "words", "1500")
        self.assertEqual(cfg.words, 1500)
        cfgmod.apply_setting(cfg, "words", "0")             # 0 allowed (off)
        cfgmod.apply_setting(cfg, "prose_style", "essay")
        self.assertEqual(cfg.prose_style, "essay")
        cfgmod.apply_setting(cfg, "enable_charts", "false")
        self.assertFalse(cfg.enable_charts)
        with self.assertRaises(ValueError):
            cfgmod.apply_setting(cfg, "prose_style", "novel")   # invalid choice

        with self.assertRaises(ValueError):
            cfgmod.apply_setting(cfg, "citations", "fancy")     # invalid choice
        with self.assertRaises(ValueError):
            cfgmod.apply_setting(cfg, "max_concurrency", "0")   # must be >= 1
        with self.assertRaises(ValueError):
            cfgmod.apply_setting(cfg, "max_retries", "-1")      # must be >= 0
        with self.assertRaises(ValueError):
            cfgmod.apply_setting(cfg, "max_parallel", "not-an-int")
        with self.assertRaises(ValueError):
            cfgmod.apply_setting(cfg, "max_parallel", "0")     # must be >= 1
        with self.assertRaises(ValueError):
            cfgmod.apply_setting(cfg, "subquestions", "-2")    # must be >= 1
        with self.assertRaises(ValueError):
            cfgmod.apply_setting(cfg, "quality_threshold", "2.0")  # out of [0,1]
        with self.assertRaises(ValueError):
            cfgmod.apply_setting(cfg, "enable_memory", "maybe")
        with self.assertRaises(ValueError):
            cfgmod.apply_setting(cfg, "mode", "bogus")
        with self.assertRaises(KeyError):
            cfgmod.apply_setting(cfg, "nope", "x")

    def test_load_coerces_bad_types(self):
        import json as _json

        path = cfgmod.config_path()
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            _json.dumps({
                "quality_threshold": "high",   # wrong type -> default kept
                "max_parallel": "lots",         # wrong type -> default kept
                "mode": "bogus",                # invalid -> default kept
                "enable_memory": "nope",         # invalid bool -> default kept
                "planner_model": "custom-model", # valid string -> passes through
            }),
            encoding="utf-8",
        )
        cfg = cfgmod.load_config()
        self.assertEqual(cfg.quality_threshold, 0.75)
        self.assertEqual(cfg.max_parallel, 2)
        self.assertEqual(cfg.mode, "auto")
        self.assertTrue(cfg.enable_memory)
        self.assertEqual(cfg.planner_model, "custom-model")


class TestModeSelection(unittest.TestCase):
    def test_explicit_modes_passthrough(self):
        self.assertEqual(choose_mode("api"), "api")
        self.assertEqual(choose_mode("subscription"), "subscription")
        self.assertEqual(choose_mode("openai"), "openai")


class TestOpenAI(unittest.TestCase):
    def test_map_openai_model(self):
        from vibe_research.backends import map_openai_model

        self.assertEqual(map_openai_model("claude-opus-4-8"), "gpt-4o")
        self.assertEqual(map_openai_model("claude-sonnet-4-6"), "gpt-4o-mini")
        self.assertEqual(map_openai_model("gpt-4o"), "gpt-4o")   # explicit passthrough
        self.assertEqual(map_openai_model("o3"), "o3")           # explicit passthrough
        self.assertEqual(map_openai_model(""), "gpt-4o")         # fallback

    def test_resolve_models(self):
        from vibe_research.backends import resolve_models

        self.assertEqual(
            resolve_models("openai", "claude-opus-4-8", "claude-sonnet-4-6"),
            ("gpt-4o", "gpt-4o-mini"),
        )
        self.assertEqual(
            resolve_models("api", "claude-opus-4-8", "claude-sonnet-4-6"),
            ("claude-opus-4-8", "claude-sonnet-4-6"),
        )

    def test_detect_available_has_openai_keys(self):
        from vibe_research.backends import detect_available

        avail = detect_available()
        self.assertIn("openai", avail)
        self.assertIn("openai_key_set", avail)

    @unittest.skipUnless(importlib.util.find_spec("openai"), "openai not installed")
    def test_openai_backend_requires_key(self):
        from vibe_research.backends import OpenAIBackend

        old = os.environ.pop("OPENAI_API_KEY", None)
        try:
            with self.assertRaises(RuntimeError):
                OpenAIBackend()
        finally:
            if old is not None:
                os.environ["OPENAI_API_KEY"] = old


class TestReports(unittest.TestCase):
    def test_save_and_list(self):
        with tempfile.TemporaryDirectory() as d:
            path = reportsmod.save_report(pathlib.Path(d), "My Topic", "# Title\n\nBody")
            self.assertTrue(path.exists())
            self.assertIn("Body", path.read_text(encoding="utf-8"))
            listed = reportsmod.list_reports(pathlib.Path(d))
            self.assertEqual(len(listed), 1)

    def test_header_added_when_missing(self):
        with tempfile.TemporaryDirectory() as d:
            path = reportsmod.save_report(pathlib.Path(d), "T", "no heading here")
            content = path.read_text(encoding="utf-8")
            self.assertTrue(content.startswith("# T"))

    def test_metadata_header(self):
        with tempfile.TemporaryDirectory() as d:
            meta = {
                "mode": "api",
                "overall_confidence": 0.82,
                "sources": ["a", "b"],
                "subquestions": ["q1", "q2", "q3"],
            }
            path = reportsmod.save_report(pathlib.Path(d), "My Topic", "# My Topic\n\nBody", meta=meta)
            content = path.read_text(encoding="utf-8")
            self.assertTrue(content.startswith("# My Topic"))
            self.assertIn("> **Generated**", content)
            self.assertIn("**confidence** 82%", content)
            self.assertIn("**sources** 2", content)
            self.assertIn("**sub-questions** 3", content)
            self.assertIn("Body", content)

    def test_save_json_sidecar(self):
        import json as _json

        with tempfile.TemporaryDirectory() as d:
            md = pathlib.Path(d) / "20260101-000000-x.md"
            md.write_text("stub", encoding="utf-8")
            jp = reportsmod.save_json(md, {"topic": "t", "findings": []})
            self.assertEqual(jp.suffix, ".json")
            self.assertEqual(jp.stem, md.stem)
            self.assertEqual(_json.loads(jp.read_text(encoding="utf-8"))["topic"], "t")


class TestMemory(unittest.TestCase):
    def test_remember_recall_roundtrip(self):
        with tempfile.TemporaryDirectory() as d:
            mem = Memory(d)
            rec = ResearchRecord(
                topic="Gut microbiota and anxiety",
                key_points=["butyrate matters"],
                sources=["https://a.com"],
                confidence=0.8,
            )
            mem.remember(rec)
            got = mem.recall("Gut microbiota and anxiety")
            self.assertIsNotNone(got)
            self.assertEqual(got.confidence, 0.8)
            self.assertTrue(got.created)  # stamped on save

    def test_related_by_overlap_excludes_self(self):
        with tempfile.TemporaryDirectory() as d:
            mem = Memory(d)
            mem.remember(ResearchRecord(topic="microbiota and depression", key_points=["serotonin gut"]))
            mem.remember(ResearchRecord(topic="stock market volatility", key_points=["vix index"]))
            related = mem.related("microbiota and anxiety disorders")
            self.assertEqual(len(related), 1)
            self.assertEqual(related[0].topic, "microbiota and depression")
            # exact-slug self is excluded
            self.assertEqual(mem.related("microbiota and depression"), [])

    def test_clear(self):
        with tempfile.TemporaryDirectory() as d:
            mem = Memory(d)
            mem.remember(ResearchRecord(topic="a longer topic one"))
            mem.remember(ResearchRecord(topic="a longer topic two"))
            self.assertEqual(mem.clear(), 2)
            self.assertEqual(mem.all(), [])

    def test_slug_collision_does_not_overwrite(self):
        with tempfile.TemporaryDirectory() as d:
            mem = Memory(d)
            mem.remember(ResearchRecord(topic="C++ vs C#", confidence=0.9))
            mem.remember(ResearchRecord(topic="C vs C", confidence=0.1))
            # Both slug to 'c-vs-c' but the hash suffix keeps them distinct.
            self.assertEqual(len(mem.all()), 2)
            self.assertEqual(mem.recall("C++ vs C#").confidence, 0.9)
            self.assertEqual(mem.recall("C vs C").confidence, 0.1)


class TestAgents(unittest.TestCase):
    def test_planner_returns_validated_plan(self):
        plan = asyncio.run(PlannerAgent(FakeBackend(), "m").plan("topic", 3))
        self.assertIsInstance(plan, ResearchPlan)
        self.assertEqual(len(plan.subquestions), 3)
        self.assertTrue(all(len(q) > 6 for q in plan.questions))

    def test_verifier_debate_aggregates(self):
        finding = Finding.build("q", "answer", ["https://a.com"])
        report = asyncio.run(VerifierAgent(FakeBackend(), "m").debate("topic", finding, votes=3))
        self.assertIsInstance(report, VerificationReport)
        self.assertAlmostEqual(report.overall_confidence, 0.9, places=2)

    def test_extract_confidence_ranges(self):
        self.assertEqual(_extract_confidence("...\nCONFIDENCE: 0.9"), 0.9)
        self.assertEqual(_extract_confidence("CONFIDENCE: 100"), 0.6)  # out of range -> default
        self.assertEqual(_extract_confidence("CONFIDENCE: 10"), 0.6)
        self.assertEqual(_extract_confidence("no hint"), 0.6)

    def test_researcher_survives_backend_error(self):
        finding = asyncio.run(
            ResearcherAgent(FailResearchBackend(), "m").research("Question?", asyncio.Semaphore(1))
        )
        self.assertEqual(finding.confidence, 0.0)
        self.assertEqual(finding.sources, [])
        self.assertIn("could not be completed", finding.answer)

    def test_verifier_survives_backend_error(self):
        finding = Finding.build("q", "a", ["https://a.com"])
        report = asyncio.run(VerifierAgent(FailVerifyBackend(), "m").debate("t", finding, votes=2))
        self.assertIsInstance(report, VerificationReport)  # abstains, does not raise

    def test_style_instructions(self):
        from vibe_research.agents import _style_instructions

        full = _style_instructions("essay", 1500, True, True, True)
        self.assertIn("essay", full.lower())
        self.assertIn("1500", full)
        self.assertIn("chart", full.lower())
        self.assertIn("mermaid", full.lower())
        minimal = _style_instructions("report", 0, False, False, False)
        self.assertNotIn("1500", minimal)
        self.assertNotIn("```chart", minimal)
        self.assertNotIn("mermaid", minimal.lower())

    def test_humanizer_keeps_rewrite_when_sources_preserved(self):
        from vibe_research.agents import HumanizerAgent

        original = "Draft body with https://example.com/a.\n\n## Confidence & Gaps\nok"
        out = asyncio.run(HumanizerAgent(FakeBackend(), "m").humanize("topic", original))
        self.assertIn("human-sounding", out)              # used the rewrite
        self.assertIn("https://example.com/a", out)       # citation preserved
        self.assertIn("Confidence & Gaps", out)

    def test_humanizer_falls_back_when_sources_lost(self):
        from vibe_research.agents import HumanizerAgent

        original = "Body with https://example.com/a and https://example.com/b.\n\n## Gaps\nok"
        out = asyncio.run(HumanizerAgent(StripUrlBackend(), "m").humanize("topic", original))
        self.assertEqual(out, original)                    # rewrite dropped URLs -> keep draft

    def test_humanizer_falls_back_when_confidence_section_lost(self):
        from vibe_research.agents import HumanizerAgent

        original = "Body https://example.com/a.\n\n## Confidence & Gaps\nok"
        out = asyncio.run(HumanizerAgent(DropSectionBackend(), "m").humanize("topic", original))
        self.assertEqual(out, original)                    # dropped honesty section -> keep draft


class TestOrchestrator(unittest.TestCase):
    def test_full_run_approve_path(self):
        backend = FakeBackend()
        events = []
        report = asyncio.run(
            run_pipeline(
                backend,
                "Test topic",
                planner_model="planner",
                worker_model="worker",
                subquestions=3,
                max_parallel=2,
                on_event=lambda k, d: events.append((k, d)),
            )
        )

        self.assertIn("Confidence & Gaps", report)
        self.assertIn("## Sources", report)  # credibility-ranked references
        self.assertIn("https://example.com/x", report)

        kinds = [k for k, _ in events]
        self.assertEqual(kinds[0], "start")
        self.assertIn("plan", kinds)
        self.assertEqual(kinds.count("finding"), 3)
        self.assertIn("debate", kinds)
        self.assertIn("critique", kinds)
        self.assertEqual(kinds[-1], "done")

        # planner model for plan/verify/edit/write, worker for research
        models_used = {m for m, _ in backend.calls}
        self.assertIn("planner", models_used)
        self.assertIn("worker", models_used)
        # research requested web search
        self.assertTrue(any(search for _, search in backend.calls))
        # done carries an overall confidence
        done = [d for k, d in events if k == "done"][0]
        self.assertGreater(done["confidence"], 0.0)

    def test_self_refining_iteration_adds_gap_research(self):
        backend = IteratingBackend()
        events = []
        asyncio.run(
            run_pipeline(
                backend,
                "Test topic",
                planner_model="planner",
                worker_model="worker",
                subquestions=3,
                max_parallel=2,
                max_iterations=2,
                on_event=lambda k, d: events.append((k, d)),
            )
        )
        kinds = [k for k, _ in events]
        # 3 initial + 1 gap-filling finding
        self.assertEqual(kinds.count("finding"), 4)
        self.assertIn("iteration", kinds)
        self.assertEqual(backend.editor_calls, 2)

    def test_memory_persisted_after_run(self):
        with tempfile.TemporaryDirectory() as d:
            mem = Memory(d)
            events = []
            asyncio.run(
                run_pipeline(
                    FakeBackend(),
                    "Persisted topic",
                    planner_model="planner",
                    worker_model="worker",
                    subquestions=3,
                    max_parallel=2,
                    enable_memory=True,
                    memory=mem,
                    on_event=lambda k, d: events.append((k, d)),
                )
            )
            record = mem.recall("Persisted topic")
            self.assertIsNotNone(record)
            self.assertGreater(len(record.subquestions), 0)
            self.assertIn("https://example.com/x", record.sources)
            self.assertTrue(any(k == "memory" and d.get("saved") for k, d in events))

    def test_humanize_stage_runs_and_can_be_disabled(self):
        ev1 = []
        asyncio.run(run_pipeline(
            FakeBackend(), "T", planner_model="p", worker_model="w",
            subquestions=3, max_parallel=2, on_event=lambda k, d: ev1.append((k, d)),
        ))
        self.assertIn("humanize", [d.get("stage") for k, d in ev1 if k == "stage"])

        ev2 = []
        asyncio.run(run_pipeline(
            FakeBackend(), "T", planner_model="p", worker_model="w",
            subquestions=3, max_parallel=2, humanize=False,
            on_event=lambda k, d: ev2.append((k, d)),
        ))
        self.assertNotIn("humanize", [d.get("stage") for k, d in ev2 if k == "stage"])

    def test_append_sources_respects_existing_references_heading(self):
        from vibe_research import pipeline

        findings = [Finding.build("q", "a", ["https://cdc.gov/x"])]
        report = "Body text.\n\n## References\n1. https://cdc.gov/x"
        out = pipeline._append_sources(report, findings, "ranked")
        self.assertNotIn("## Sources", out)   # don't append a duplicate list
        self.assertEqual(out, report)

    def test_disagreements_section_surfaced(self):
        report = asyncio.run(run_pipeline(
            ContradictionBackend(), "T", planner_model="p", worker_model="w",
            subquestions=3, max_parallel=2,
        ))
        self.assertIn("## Disagreements", report)

    def test_per_stage_model_override(self):
        backend = FakeBackend()
        asyncio.run(run_pipeline(
            backend, "T", planner_model="planner", worker_model="worker",
            subquestions=3, max_parallel=2,
            verifier_model="verif-x", writer_model="writer-x",
        ))
        models = {m for m, _ in backend.calls}
        self.assertIn("verif-x", models)
        self.assertIn("writer-x", models)

    def test_domain_block_filter_applied(self):
        events = []
        asyncio.run(run_pipeline(
            BlockedSourceBackend(), "T", planner_model="p", worker_model="w",
            subquestions=3, max_parallel=2, block_domains=["reddit.com"],
            on_event=lambda k, d: events.append((k, d)),
        ))
        result = [d["result"] for k, d in events if k == "done"][0]
        self.assertNotIn("https://reddit.com/r/x", result["sources"])
        self.assertIn("https://cdc.gov/y", result["sources"])

    def test_run_survives_transient_backend_failures(self):
        # A raising researcher OR a raising verifier must degrade, not crash.
        for backend in (FailResearchBackend(), FailVerifyBackend()):
            events = []
            report = asyncio.run(
                run_pipeline(
                    backend,
                    "Resilience topic",
                    planner_model="planner",
                    worker_model="worker",
                    subquestions=3,
                    max_parallel=2,
                    on_event=lambda k, d: events.append((k, d)),
                )
            )
            kinds = [k for k, _ in events]
            self.assertEqual(kinds[-1], "done")          # completed cleanly
            self.assertEqual(kinds.count("finding"), 3)
            self.assertIn("Confidence & Gaps", report)


class TestReliability(unittest.TestCase):
    def test_retry_then_success(self):
        b = FlakyBackend(fail_times=2)
        text, _ = asyncio.run(b.complete("p", "m"))
        self.assertEqual(text, "ok")
        self.assertEqual(b.retries, 2)
        self.assertEqual(b.usage()["calls"], 1)

    def test_gives_up_after_max_retries(self):
        b = FlakyBackend(fail_times=99)
        b.max_retries = 2
        with self.assertRaises(ConnectionError):
            asyncio.run(b.complete("p", "m"))
        self.assertEqual(b.failures, 1)
        self.assertEqual(b.retries, 2)

    def test_non_retryable_fails_fast(self):
        b = AuthErrorBackend()
        with self.assertRaises(RuntimeError):
            asyncio.run(b.complete("p", "m"))
        self.assertEqual(b.attempts, 1)   # 401 is not retried

    def test_per_call_timeout(self):
        b = HangBackend()
        with self.assertRaises((asyncio.TimeoutError, TimeoutError)):
            asyncio.run(b.complete("p", "m"))
        self.assertEqual(b.failures, 1)

    def test_concurrency_cap(self):
        b = ThrottleBackend()

        async def go():
            await asyncio.gather(*(b.complete("p", "m") for _ in range(8)))

        asyncio.run(go())
        self.assertLessEqual(b.max_active, 2)
        self.assertEqual(b.calls, 8)

    def test_usage_line(self):
        b = FlakyBackend(fail_times=1)
        asyncio.run(b.complete("p", "m"))
        self.assertIn("calls", b.usage_line())
        self.assertIn("retries", b.usage_line())

    def test_debug_trace_written(self):
        import json as _json

        b = FlakyBackend(fail_times=0)
        with tempfile.TemporaryDirectory() as d:
            p = pathlib.Path(d) / "trace.jsonl"
            b.configure(debug_path=p)
            asyncio.run(b.complete("hello prompt", "m"))
            lines = p.read_text(encoding="utf-8").strip().splitlines()
            self.assertEqual(len(lines), 1)
            rec = _json.loads(lines[0])
            self.assertEqual(rec["model"], "m")
            self.assertIn("prompt_head", rec)


class TestEnrich(unittest.TestCase):
    def test_score_source_tiers(self):
        tier = lambda u: enrich.score_source(u)["tier"]
        self.assertEqual(tier("https://cdc.gov/x"), "high")
        self.assertEqual(tier("https://mit.edu/y"), "high")
        self.assertEqual(tier("https://nature.com/a"), "high")
        self.assertEqual(tier("https://bbc.com/n"), "medium-high")
        self.assertEqual(tier("https://redcross.org/z"), "medium")
        self.assertEqual(tier("https://reddit.com/r/x"), "low")
        self.assertEqual(tier("https://randomsite.net/x"), "medium")

    def test_score_source_matches_host_not_path(self):
        # An authoritative domain in the PATH/QUERY must not elevate a junk host.
        self.assertEqual(enrich.score_source("http://evil.example/r?to=https://cdc.gov")["tier"], "medium")
        self.assertEqual(enrich.score_source("http://myblog.example/nature.com-is-great")["tier"], "medium")
        # A genuinely authoritative host is still scored high.
        self.assertEqual(enrich.score_source("https://pubmed.ncbi.nlm.nih.gov/123")["tier"], "high")

    def test_rank_sources_orders_by_credibility(self):
        ranked = enrich.rank_sources(
            ["https://reddit.com/x", "https://cdc.gov/y", "https://bbc.com/z"]
        )
        self.assertEqual(enrich.domain_of(ranked[0]["url"]), "cdc.gov")
        self.assertEqual(enrich.domain_of(ranked[-1]["url"]), "reddit.com")

    def test_sources_section(self):
        section = enrich.sources_section(["https://cdc.gov/a"])
        self.assertIn("## Sources", section)
        self.assertIn("1. https://cdc.gov/a", section)
        self.assertIn("primary", section)
        self.assertEqual(enrich.sources_section([]), "")

    def test_filter_sources(self):
        urls = ["https://a.gov/x", "https://reddit.com/y", "https://b.edu/z"]
        self.assertEqual(
            enrich.filter_sources(urls, only=["gov", "edu"]),
            ["https://a.gov/x", "https://b.edu/z"],
        )
        self.assertEqual(
            enrich.filter_sources(urls, block=["reddit.com"]),
            ["https://a.gov/x", "https://b.edu/z"],
        )

    def test_disagreements_section(self):
        self.assertEqual(enrich.disagreements_section([]), "")
        section = enrich.disagreements_section(["A vs B", "A vs B", "C conflict"])
        self.assertIn("## Disagreements", section)
        self.assertEqual(section.count("\n- "), 2)  # deduped to two bullets

    def test_credibility_summary(self):
        summary = enrich.credibility_summary(["https://cdc.gov", "https://reddit.com/x"])
        self.assertIn("high", summary)
        self.assertIn("low", summary)


class TestVisuals(unittest.TestCase):
    def test_count_words_ignores_code(self):
        self.assertEqual(visuals.count_words("one two three"), 3)
        self.assertEqual(visuals.count_words("a b\n```\nx y z q\n```\nc"), 3)

    def test_parse_chart_spec(self):
        spec = visuals.parse_chart_spec(
            '{"type":"bar","title":"T","labels":["a","b"],"series":[{"name":"s","data":[1,2]}]}'
        )
        self.assertEqual(spec["type"], "bar")
        self.assertEqual(spec["series"][0]["data"], [1.0, 2.0])
        # single-series via 'values'
        spec2 = visuals.parse_chart_spec('{"labels":["a"],"values":[5]}')
        self.assertEqual(spec2["series"][0]["data"], [5.0])
        self.assertIsNone(visuals.parse_chart_spec("not json"))
        self.assertIsNone(visuals.parse_chart_spec('{"labels":["a"]}'))  # no data

    @unittest.skipUnless(visuals.charts_available(), "matplotlib not installed")
    def test_render_chart_writes_png(self):
        with tempfile.TemporaryDirectory() as d:
            out = pathlib.Path(d) / "c.png"
            spec = visuals.parse_chart_spec('{"type":"line","labels":["a","b","c"],"values":[1,2,3]}')
            self.assertIsNotNone(visuals.render_chart(spec, out))
            self.assertGreater(out.stat().st_size, 0)

    @unittest.skipUnless(visuals.charts_available(), "matplotlib not installed")
    def test_render_report_charts_replaces_block(self):
        with tempfile.TemporaryDirectory() as d:
            report = 'Intro\n\n```chart\n{"type":"bar","labels":["a"],"values":[3]}\n```\n\nEnd'
            out = visuals.render_report_charts(report, d, "rep")
            self.assertIn("![", out)
            self.assertIn(".png", out)
            self.assertNotIn("```chart", out)
            self.assertEqual(len([f for f in os.listdir(d) if f.endswith(".png")]), 1)

    def test_render_report_charts_table_fallback(self):
        orig = visuals.render_chart
        visuals.render_chart = lambda spec, out: None   # simulate no matplotlib
        try:
            report = '```chart\n{"type":"bar","labels":["a","b"],"values":[1,2]}\n```'
            out = visuals.render_report_charts(report, ".", "x")
            self.assertIn("| Category |", out)   # table fallback, no data lost
            self.assertNotIn("```chart", out)
        finally:
            visuals.render_chart = orig


class TestCLI(unittest.TestCase):
    def _cfg_from(self, **kwargs):
        import argparse

        from vibe_research import cli

        cfg = cfgmod.default_config()
        cli._cfg_from_args(cfg, argparse.Namespace(**kwargs))
        return cfg

    def test_depth_preset(self):
        cfg = self._cfg_from(depth="deep")
        self.assertEqual((cfg.subquestions, cfg.verifier_votes, cfg.max_iterations), (8, 3, 3))

    def test_depth_preset_explicit_override_wins(self):
        cfg = self._cfg_from(depth="quick", subquestions=6)
        self.assertEqual(cfg.subquestions, 6)       # explicit beats preset
        self.assertEqual(cfg.verifier_votes, 1)     # rest still from 'quick'

    def test_pages_maps_to_words(self):
        self.assertEqual(self._cfg_from(pages=3).words, 1500)
        self.assertEqual(self._cfg_from(words=800).words, 800)

    def test_style_flag(self):
        self.assertEqual(self._cfg_from(style="essay").prose_style, "essay")

    def test_cost_estimate(self):
        from vibe_research import cli

        self.assertEqual(cli._cost_estimate({}), "")
        line = cli._cost_estimate({"input_tokens": 1_000_000, "output_tokens": 1_000_000})
        self.assertIn("$", line)


class TestExport(unittest.TestCase):
    def test_markdown_to_html(self):
        html = exportmod.markdown_to_html(
            "# T\n\n- a\n- b\n\n[x](https://e.com)", title="My<>Title"
        )
        self.assertIn("<!doctype html>", html.lower())
        self.assertIn("<h1>", html)
        self.assertIn("https://e.com", html)
        self.assertNotIn("My<>Title", html)  # title HTML-escaped

    def test_html_path_for(self):
        self.assertEqual(exportmod.html_path_for("a/b.md").suffix, ".html")

    def test_docx_path_for(self):
        self.assertEqual(exportmod.docx_path_for("a/b.md").suffix, ".docx")

    def test_html_renders_mermaid(self):
        html = exportmod.markdown_to_html("# T\n\n```mermaid\nflowchart TD\n  A-->B\n```\n")
        self.assertIn('class="mermaid"', html)
        self.assertIn("mermaid.esm.min.mjs", html)
        self.assertIn("A--&gt;B" if False else "A-->B", html)  # arrow preserved, unescaped

    def test_resolve_images_makes_local_absolute_keeps_remote(self):
        out = exportmod._resolve_images("![c](chart-1.png) ![w](https://x.com/i.png)", "/base/dir")
        norm = out.replace("\\", "/")
        self.assertIn("/base/dir", norm)               # relative -> absolute
        self.assertIn("https://x.com/i.png", out)      # remote untouched

    @unittest.skipUnless(exportmod.docx_available(), "python-docx not installed")
    def test_markdown_to_docx_writes_file(self):
        with tempfile.TemporaryDirectory() as d:
            out = pathlib.Path(d) / "r.docx"
            md = "# T\n\nA **bold** and *em* point with [link](https://e.com).\n\n- one\n- two"
            path = exportmod.markdown_to_docx(md, out, title="T")
            self.assertTrue(path.exists())
            self.assertGreater(path.stat().st_size, 0)

    @unittest.skipUnless(exportmod.docx_available(), "python-docx not installed")
    def test_markdown_to_docx_renders_table_and_heading_markup(self):
        import docx

        with tempfile.TemporaryDirectory() as d:
            out = pathlib.Path(d) / "t.docx"
            md = "## **Key** findings\n\n| Model | Cost |\n| --- | --- |\n| Opus | High |\n"
            exportmod.markdown_to_docx(md, out, title="T")
            doc = docx.Document(str(out))
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
            self.assertIn("Key findings", headings)          # inline markup rendered
            self.assertNotIn("**Key** findings", headings)   # not raw source
            self.assertGreaterEqual(len(doc.tables), 1)      # table preserved
            cells = [c.text for t in doc.tables for row in t.rows for c in row.cells]
            self.assertIn("Opus", cells)
            self.assertIn("Cost", cells)

    @unittest.skipUnless(exportmod.pdf_available(), "fpdf2 not installed")
    def test_markdown_to_pdf_writes_valid_file(self):
        with tempfile.TemporaryDirectory() as d:
            out = pathlib.Path(d) / "r.pdf"
            md = (
                "# Title\n\nA **bold** point → arrow, accents café/naïve.\n\n"
                "- one\n- two\n\n## Confidence & Gaps\nok"
            )
            path = exportmod.markdown_to_pdf(md, out, title="Title")
            self.assertTrue(path.exists())
            self.assertGreater(path.stat().st_size, 0)
            with open(path, "rb") as fh:
                self.assertTrue(fh.read(5).startswith(b"%PDF"))

    def test_pdf_path_for(self):
        self.assertEqual(exportmod.pdf_path_for(pathlib.Path("/x/y.md")).name, "y.pdf")
        self.assertEqual(exportmod.pdf_path_for("a/b/c.md").suffix, ".pdf")


class TestCliHeadless(unittest.TestCase):
    """Exercise the CLI headless glue end-to-end through the real entry point."""

    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self._old = os.environ.get("XDG_CONFIG_HOME")
        os.environ["XDG_CONFIG_HOME"] = self._tmp.name

    def tearDown(self):
        if self._old is None:
            os.environ.pop("XDG_CONFIG_HOME", None)
        else:
            os.environ["XDG_CONFIG_HOME"] = self._old
        self._tmp.cleanup()

    def test_headless_run_writes_md_and_optional_pdf(self):
        from vibe_research import backends as backendsmod
        from vibe_research import cli

        want_pdf = exportmod.pdf_available()
        with tempfile.TemporaryDirectory() as d:
            argv = ["run", "Headless integration topic", "--no-tui",
                    "--output-dir", d, "--no-memory"]
            if want_pdf:
                argv.append("--pdf")

            original = backendsmod.get_backend
            backendsmod.get_backend = lambda mode="auto": FakeBackend()
            try:
                rc = cli.main(argv)
            finally:
                backendsmod.get_backend = original

            self.assertEqual(rc, 0)
            mds = list(pathlib.Path(d).glob("*.md"))
            self.assertEqual(len(mds), 1)
            self.assertIn("Confidence & Gaps", mds[0].read_text(encoding="utf-8"))
            if want_pdf:
                self.assertEqual(len(list(pathlib.Path(d).glob("*.pdf"))), 1)


class TestTUI(unittest.TestCase):
    @unittest.skipUnless(importlib.util.find_spec("textual"), "textual not installed")
    def test_app_constructs_with_new_actions_and_bindings(self):
        from vibe_research.tui import VibeResearchApp

        app = VibeResearchApp(cfgmod.default_config(), "topic")
        for action in ("action_export_all", "action_open_report", "action_export_pdf"):
            self.assertTrue(callable(getattr(app, action, None)), action)
        keys = {b.key for b in VibeResearchApp.BINDINGS}
        self.assertIn("ctrl+e", keys)  # export all
        self.assertIn("ctrl+o", keys)  # open report

    @unittest.skipUnless(importlib.util.find_spec("textual"), "textual not installed")
    def test_usage_str_reads_backend_tokens(self):
        from vibe_research.tui import VibeResearchApp

        app = VibeResearchApp(cfgmod.default_config())
        self.assertEqual(app._usage_str(), "")   # no backend yet -> empty

        backend = FlakyBackend(fail_times=0)
        backend.calls, backend.input_tokens, backend.output_tokens = 3, 1200, 800
        app._backend = backend
        usage = app._usage_str()
        self.assertIn("calls", usage)
        self.assertIn("2.0k tok", usage)   # 1200 + 800 tokens


if __name__ == "__main__":
    unittest.main()
